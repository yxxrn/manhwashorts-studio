"""Source material ingestion (PRD FR-02).

Accepts text, documents (TXT/MD/PDF/DOCX), and images. Every asset must carry
a rights declaration before it can reach publication; see
``app.services.policy``.

Deliberately absent: any form of remote fetching or scraping. Material only
enters the system through explicit user upload, per PRD section 8.
"""

from __future__ import annotations

import io
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageFilter, UnidentifiedImageError

from app.config import settings
from app.constants import AssetType, LicenseType, RightsStatus
from app.services import storage


class IngestError(ValueError):
    """Raised when submitted material cannot be accepted."""


@dataclass
class RightsDeclaration:
    """User's assertion about how they are allowed to use a piece of material."""

    source_name: str = ""
    rights_owner: str = ""
    license_type: str = LicenseType.UNKNOWN
    permission_reference: str = ""
    permission_date: str = ""
    usage_limits: str = ""
    attribution: str = ""
    declared: bool = False

    @property
    def status(self) -> str:
        """Declared requires an owner and a concrete licence basis."""
        if not self.declared:
            return RightsStatus.UNDECLARED
        if self.license_type == LicenseType.UNKNOWN or not self.rights_owner.strip():
            return RightsStatus.UNDECLARED
        return RightsStatus.DECLARED


@dataclass
class IngestedAsset:
    """Normalised result of ingesting one piece of material."""

    type: str
    original_filename: str
    mime_type: str
    storage_key: str
    size_bytes: int
    checksum: str
    extracted_text: str = ""
    width: int = 0
    height: int = 0
    source_family: str = ""
    panel_bbox: dict = None
    panel_quality: dict = None
    panel_decision: str = "accept"
    audio_duration: float = 0.0
    original_checksum: str = ""
    original_width: int = 0
    original_height: int = 0
    source_bounds: tuple[int, int, int, int] = (0, 0, 0, 0)
    strip_order: int = 0
    region_order: int = 0
    trim_classification: str = "unsliced"
    coverage_map_hash: str = ""


_SLICE_SUFFIX = re.compile(r"(?:[_-]p?\d+)$", re.IGNORECASE)


def derive_source_family(filename: str) -> str:
    """Stable logical family for an uploaded page and its sliced panels."""
    value = str(filename or "").replace("\\", "/").strip("/")
    if not value:
        return "unknown"
    path = Path(value)
    stem = (_SLICE_SUFFIX.sub("", path.stem).rstrip("_-") or path.stem or "unknown")
    parent = "/".join(part for part in path.parts[:-1] if part not in {"", "."})
    return f"{parent}/{stem}" if parent else stem


# --- text extraction -------------------------------------------------------

_WS = re.compile(r"[ \t]+")
_BLANKS = re.compile(r"\n{3,}")


def normalise_text(raw: str) -> str:
    """Collapse whitespace while preserving paragraph breaks."""
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return _BLANKS.sub("\n\n", text).strip()


def extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        parts = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:  # pypdf raises a wide variety of errors
        raise IngestError(f"could not read PDF: {exc}") from exc
    text = normalise_text("\n\n".join(parts))
    if not text:
        raise IngestError(
            "no selectable text found in PDF (scanned images are not supported; "
            "paste the recap as text instead)"
        )
    return text


def extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise IngestError(f"could not read DOCX: {exc}") from exc
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(c.text for c in row.cells))
    text = normalise_text("\n".join(paragraphs))
    if not text:
        raise IngestError("DOCX contained no text")
    return text


def extract_document(filename: str, mime_type: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf" or mime_type == "application/pdf":
        return extract_pdf(data)
    if suffix == ".docx" or "wordprocessingml" in mime_type:
        return extract_docx(data)
    if suffix in {".txt", ".md", ".markdown"} or mime_type.startswith("text/"):
        try:
            return normalise_text(data.decode("utf-8"))
        except UnicodeDecodeError:
            return normalise_text(data.decode("latin-1", errors="replace"))
    raise IngestError(f"unsupported document type: {filename} ({mime_type})")


# --- validation ------------------------------------------------------------


def _check_size(data: bytes) -> None:
    if len(data) == 0:
        raise IngestError("file is empty")
    if len(data) > settings.max_upload_bytes:
        raise IngestError(
            f"file exceeds {settings.max_upload_mb} MB limit "
            f"({len(data) / 1024 / 1024:.1f} MB)"
        )


def _panel_quality(data: bytes) -> dict:
    """Cheap CPU-only panel quality metadata; rejects blank connector strips."""
    with Image.open(io.BytesIO(data)) as image:
        gray = image.convert("L").resize((96, 96), Image.Resampling.BILINEAR)
        pixels = list(gray.getdata())
        blank = sum(pixel >= 245 or pixel <= 10 for pixel in pixels) / max(1, len(pixels))
        edges = gray.filter(ImageFilter.FIND_EDGES)
        edge_density = sum(pixel > 35 for pixel in edges.getdata()) / max(1, len(pixels))
        width, height = image.size
    decision = "reject" if blank >= 0.94 or (edge_density < 0.008 and blank >= 0.82) else "accept"
    return {
        "bbox": {"x": 0, "y": 0, "width": width, "height": height},
        "content_coverage": round(1.0 - blank, 4),
        "blank_ratio": round(blank, 4),
        "edge_density": round(edge_density, 4),
        "aspect_ratio": round(width / max(1, height), 4),
        "ocr_regions": [],
        "decision": decision,
        "reason": "blank/gutter-like candidate" if decision == "reject" else "content-bearing panel",
    }


def _sniff_image(data: bytes) -> tuple[str, int, int]:
    """Verify the bytes really are an image and return (mime, w, h).

    Trusting a client-supplied Content-Type is not enough: Pillow re-parses
    the payload so a renamed executable cannot slip into the pipeline.
    """
    try:
        with Image.open(io.BytesIO(data)) as img:
            img.verify()
        with Image.open(io.BytesIO(data)) as img:
            width, height = img.size
            fmt = (img.format or "").upper()
    except (UnidentifiedImageError, OSError, ValueError) as exc:
        raise IngestError("file is not a valid image") from exc

    mime = {"JPEG": "image/jpeg", "PNG": "image/png", "WEBP": "image/webp"}.get(fmt, "")
    if mime not in settings.allowed_image_types:
        raise IngestError(f"unsupported image format: {fmt or 'unknown'} (use JPEG, PNG, or WebP)")
    if width < 200 or height < 200:
        raise IngestError(f"image too small ({width}x{height}); minimum 200x200")
    return mime, width, height


# --- ingest entry points ---------------------------------------------------


def ingest_text(project_id: str, text: str, title: str = "notes.txt") -> IngestedAsset:
    """Store a pasted recap or note as a text asset."""
    cleaned = normalise_text(text)
    if len(cleaned) < 40:
        raise IngestError("text is too short to summarise (minimum 40 characters)")
    data = cleaned.encode("utf-8")
    _check_size(data)
    obj = storage.put_bytes(f"projects/{project_id}/text", title, data)
    return IngestedAsset(
        type=AssetType.TEXT,
        original_filename=title,
        mime_type="text/plain",
        storage_key=obj.storage_key,
        size_bytes=obj.size_bytes,
        checksum=obj.checksum,
        extracted_text=cleaned,
    )


def ingest_document(project_id: str, filename: str, mime_type: str, data: bytes) -> IngestedAsset:
    _check_size(data)
    text = extract_document(filename, mime_type, data)
    obj = storage.put_bytes(f"projects/{project_id}/docs", filename, data)
    return IngestedAsset(
        type=AssetType.DOCUMENT,
        original_filename=filename,
        mime_type=mime_type or "application/octet-stream",
        storage_key=obj.storage_key,
        size_bytes=obj.size_bytes,
        checksum=obj.checksum,
        extracted_text=text,
    )


def ingest_audio(project_id: str, filename: str, mime_type: str, data: bytes) -> IngestedAsset:
    """Store a user-provided ambience/music asset after media validation."""
    _check_size(data)
    suffix = Path(filename).suffix.lower()
    if suffix not in {".wav", ".mp3", ".ogg", ".m4a", ".flac"}:
        raise IngestError("unsupported audio format; use WAV, MP3, OGG, M4A, or FLAC")
    obj = storage.put_bytes(f"projects/{project_id}/audio-assets", filename, data)
    duration = 0.0
    try:
        result = subprocess.run(
            [settings.ffprobe_bin, "-v", "error", "-show_entries", "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", str(storage.path_for(obj.storage_key))],
            capture_output=True, text=True, timeout=60, check=True,
        )
        duration = float(result.stdout.strip() or 0.0)
    except (OSError, subprocess.CalledProcessError, subprocess.TimeoutExpired, ValueError) as exc:
        storage.delete(obj.storage_key)
        raise IngestError(f"audio validation failed: {type(exc).__name__}") from exc
    if duration <= 0.1:
        storage.delete(obj.storage_key)
        raise IngestError("audio has no usable duration")
    return IngestedAsset(
        type=AssetType.MUSIC,
        original_filename=filename,
        mime_type=mime_type or "audio/wav",
        storage_key=obj.storage_key,
        size_bytes=obj.size_bytes,
        checksum=obj.checksum,
        audio_duration=round(duration, 3),
    )


def ingest_image(project_id: str, filename: str, data: bytes) -> IngestedAsset:
    _check_size(data)
    mime, width, height = _sniff_image(data)
    obj = storage.put_bytes(f"projects/{project_id}/images", filename, data)
    quality = _panel_quality(data)
    return IngestedAsset(
        type=AssetType.IMAGE,
        original_filename=filename,
        mime_type=mime,
        storage_key=obj.storage_key,
        size_bytes=obj.size_bytes,
        checksum=obj.checksum,
        width=width,
        height=height,
        original_checksum=obj.checksum,
        original_width=width,
        original_height=height,
        source_bounds=(0, 0, width, height),
        strip_order=0,
        region_order=0,
        trim_classification="unsliced",
        coverage_map_hash="",
        source_family=derive_source_family(filename),
        panel_bbox=quality["bbox"], panel_quality=quality, panel_decision=quality["decision"],
    )


def ingest_image_parts(project_id: str, filename: str, data: bytes) -> list[IngestedAsset]:
    """Ingest an image, splitting a tall webtoon strip into scene-sized pieces.

    A webtoon page is one long vertical strip. Forcing it into a single 9:16
    frame keeps under a third of the page, so a whole story beat can vanish.
    Tall images are therefore sliced (see ``app.services.strips``) and each piece
    is stored as its own asset, preserving reading order.

    Always returns at least one asset. Anything not tall enough to slice comes
    back as a single item holding the original bytes, so callers need no special
    case. If slicing fails for any reason the original image is kept — losing a
    panel would be worse than an unsliced one.
    """
    _check_size(data)
    # Validate first: never slice bytes that are not a real image.
    _sniff_image(data)

    from app.services import strips

    try:
        pieces = strips.slice_strip(data)
    except Exception:  # noqa: BLE001 - never lose a panel over a slicing bug
        pieces = []

    if len(pieces) <= 1:
        return [ingest_image(project_id, filename, data)]

    stem = Path(filename).stem or "panel"
    suffix = Path(filename).suffix or ".jpg"
    assets: list[IngestedAsset] = []
    for number, piece in enumerate(pieces, start=1):
        # Zero-padded so lexical order matches reading order.
        part_name = f"{stem}_p{number:02d}{suffix}"
        mime, width, height = _sniff_image(piece.data)
        obj = storage.put_bytes(f"projects/{project_id}/images", part_name, piece.data)
        quality = _panel_quality(piece.data)
        quality["bbox"] = {"x": 0, "y": piece.top, "width": width, "height": height}
        assets.append(
            IngestedAsset(
                type=AssetType.IMAGE,
                original_filename=part_name,
                mime_type=mime,
                storage_key=obj.storage_key,
                size_bytes=obj.size_bytes,
                checksum=obj.checksum,
                width=width,
                height=height,
                source_family=derive_source_family(filename),
                panel_bbox=quality["bbox"],
                panel_quality=quality,
                panel_decision=quality["decision"],
                audio_duration=0.0,
                original_checksum=piece.original_checksum,
                original_width=piece.original_width or width,
                original_height=piece.original_height or height,
                source_bounds=piece.source_bounds or (0, piece.top, width, piece.bottom),
                strip_order=piece.strip_order,
                region_order=piece.region_order if piece.region_order >= 0 else number - 1,
                trim_classification=piece.trim_classification,
                coverage_map_hash=piece.coverage_map_hash,
            )
        )
    return assets


def _is_image(suffix: str, mime_type: str) -> bool:
    return suffix in {".jpg", ".jpeg", ".png", ".webp"} or mime_type.startswith("image/")


def _is_audio(suffix: str, mime_type: str) -> bool:
    return suffix in {".wav", ".mp3", ".ogg", ".m4a", ".flac"} or mime_type.startswith("audio/")


def _is_document(suffix: str, mime_type: str) -> bool:
    return suffix in {".txt", ".md", ".markdown", ".pdf", ".docx"} or (
        mime_type in settings.allowed_doc_types
    )


def _unsupported(filename: str) -> IngestError:
    return IngestError(
        f"unsupported file type: {filename}. Accepted: JPG, PNG, WebP, WAV, MP3, OGG, M4A, FLAC, TXT, MD, PDF, DOCX"
    )


def ingest_upload(project_id: str, filename: str, mime_type: str, data: bytes) -> IngestedAsset:
    """Dispatch on declared type, then verify against real content.

    Single-asset form, kept for callers that want the original image untouched.
    Use :func:`ingest_upload_parts` to have tall strips sliced.
    """
    suffix = Path(filename).suffix.lower()
    if _is_image(suffix, mime_type):
        return ingest_image(project_id, filename, data)
    if _is_audio(suffix, mime_type):
        return ingest_audio(project_id, filename, mime_type, data)
    if _is_document(suffix, mime_type):
        return ingest_document(project_id, filename, mime_type, data)
    raise _unsupported(filename)


def ingest_upload_parts(
    project_id: str, filename: str, mime_type: str, data: bytes
) -> list[IngestedAsset]:
    """Like :func:`ingest_upload`, but one tall image may yield several assets.

    This is what the upload endpoint uses: a webtoon page is a long strip, and
    splitting it into consecutive scenes preserves the beats that a single 9:16
    crop would discard. Documents and normal-shaped images return one asset.
    """
    suffix = Path(filename).suffix.lower()
    if _is_image(suffix, mime_type):
        return ingest_image_parts(project_id, filename, data)
    if _is_audio(suffix, mime_type):
        return [ingest_audio(project_id, filename, mime_type, data)]
    if _is_document(suffix, mime_type):
        return [ingest_document(project_id, filename, mime_type, data)]
    raise _unsupported(filename)
